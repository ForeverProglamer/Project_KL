import docx


# Пример дампа памяти
# memory = [[], [['0', '000000000000', '111111', '000000111111', 'RG1:= 000000000000RG2:= 111111RG3:= 000000111111'],
#                ['1', '000000111111', '011111', '000001111110', 'RG1: = RG1 + RG3RG3: = l(RG3).0RG2: = 0.r(RG2)'],
#                ['2', '000010111101', '001111', '000011111100', 'RG1: = RG1 + RG3RG3: = l(RG3).0RG2: = 0.r(RG2)'],
#                ['3', '000110111001', '000111', '000111111000', 'RG1: = RG1 + RG3RG3: = l(RG3).0RG2: = 0.r(RG2)'],
#                ['4', '001110110001', '000011', '001111110000', 'RG1: = RG1 + RG3RG3: = l(RG3).0RG2: = 0.r(RG2)'],
#                ['5', '011110100001', '000001', '011111100000', 'RG1: = RG1 + RG3RG3: = l(RG3).0RG2: = 0.r(RG2)'],
#                ['6', '111110000001', '000000', '111111000000', 'RG1: = RG1 + RG3RG3: = l(RG3).0RG2: = 0.r(RG2)']], [],
#           []]
#
# oper_num = 1


def table_save(memory, oper_num):
    """ в функцию подаем дамп памяти и номер операции """
    doc = docx.Document()
    rows = len(memory[oper_num]) + 1  # +1 строка уходит на название столбиков
    cols = len(memory[oper_num][0])

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    if cols == 5:
        doc.tables[0].cell(0, 0).text = '№'
        doc.tables[0].cell(0, 1).text = 'RG1'
        doc.tables[0].cell(0, 2).text = 'RG2'
        doc.tables[0].cell(0, 3).text = 'RG3'
        doc.tables[0].cell(0, 4).text = 'Мікрооперації'

    elif cols == 6:
        doc.tables[0].cell(0, 0).text = '№'
        doc.tables[0].cell(0, 1).text = 'RG1'
        doc.tables[0].cell(0, 2).text = 'RG2'
        doc.tables[0].cell(0, 3).text = 'RG3'
        doc.tables[0].cell(0, 4).text = 'CT'
        doc.tables[0].cell(0, 5).text = 'Мікрооперації'

    for row in range(rows - 1):
        for col in range(cols):
            doc.tables[0].cell(row + 1, col).text = memory[oper_num][row][col]

    doc.save('your table.docx')
